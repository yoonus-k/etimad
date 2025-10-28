"""
Document Processor Module
Extracts text from various document formats (PDF, Excel, Word, Images)
"""

import io
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = ['.pdf', '.xlsx', '.xls', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.txt']
    
    def process_folder(self, folder_path) -> Dict[str, any]:
        """
        Process all documents in a folder
        
        Args:
            folder_path: Path to folder containing tender documents (str or Path)
            
        Returns:
            Dict with extracted content organized by file type
        """
        # Convert to Path object if string
        if isinstance(folder_path, str):
            folder_path = Path(folder_path)
        
        if not folder_path.exists():
            logger.error(f"Folder not found: {folder_path}")
            return {}
        
        logger.info(f"📂 Processing folder: {folder_path}")
        
        result = {
            'pdfs': [],
            'excel_files': [],
            'word_docs': [],
            'images': [],
            'text_files': [],
            'total_files': 0,
            'total_text_length': 0,
            'errors': []
        }
        
        # Find all files
        files = []
        for ext in self.supported_formats:
            files.extend(list(folder_path.glob(f'*{ext}')))
            files.extend(list(folder_path.glob(f'*{ext.upper()}')))
        
        result['total_files'] = len(files)
        logger.info(f"Found {len(files)} files to process")
        
        # Process each file
        for file_path in files:
            try:
                file_ext = file_path.suffix.lower()
                
                if file_ext == '.pdf':
                    content = self._process_pdf(file_path)
                    if content:
                        result['pdfs'].append({
                            'filename': file_path.name,
                            'content': content,
                            'length': len(content)
                        })
                        result['total_text_length'] += len(content)
                
                elif file_ext in ['.xlsx', '.xls']:
                    content = self._process_excel(file_path)
                    if content:
                        result['excel_files'].append({
                            'filename': file_path.name,
                            'content': content,
                            'length': len(str(content))
                        })
                        result['total_text_length'] += len(str(content))
                
                elif file_ext in ['.docx', '.doc']:
                    content = self._process_word(file_path)
                    if content:
                        result['word_docs'].append({
                            'filename': file_path.name,
                            'content': content,
                            'length': len(content)
                        })
                        result['total_text_length'] += len(content)
                
                elif file_ext in ['.png', '.jpg', '.jpeg']:
                    # Will be processed by OCR module
                    result['images'].append({
                        'filename': file_path.name,
                        'path': str(file_path)
                    })
                
                elif file_ext == '.txt':
                    content = self._process_text(file_path)
                    if content:
                        result['text_files'].append({
                            'filename': file_path.name,
                            'content': content,
                            'length': len(content)
                        })
                        result['total_text_length'] += len(content)
                
                logger.info(f"✅ Processed: {file_path.name}")
                
            except Exception as e:
                error_msg = f"Error processing {file_path.name}: {str(e)}"
                logger.error(f"❌ {error_msg}")
                result['errors'].append(error_msg)
        
        logger.info(f"✅ Processing complete: {result['total_text_length']} characters extracted")
        return result
    
    def _process_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            logger.warning("pdfplumber not installed, trying PyPDF2")
            return self._process_pdf_pypdf2(file_path)
        
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def _process_pdf_pypdf2(self, file_path: Path) -> str:
        """Fallback PDF processing with PyPDF2"""
        try:
            from PyPDF2 import PdfReader
            
            text_parts = []
            reader = PdfReader(str(file_path))
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
            
            return "\n\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")
            return ""
    
    def _process_excel(self, file_path: Path) -> Dict:
        """
        Extract data from Excel file
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dict with sheet data
        """
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            result = {
                'sheets': {},
                'text_summary': []
            }
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to text representation
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False)
                
                result['sheets'][sheet_name] = {
                    'rows': len(df),
                    'columns': list(df.columns),
                    'data': df.to_dict('records')[:100],  # Limit to first 100 rows
                    'text': sheet_text
                }
                
                result['text_summary'].append(sheet_text)
            
            return result
        
        except Exception as e:
            logger.error(f"Error processing Excel: {e}")
            return {}
    
    def _process_word(self, file_path: Path) -> str:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            return ""
    
    def _process_text(self, file_path: Path) -> str:
        """
        Read plain text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='cp1256') as f:  # Arabic Windows encoding
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text file: {e}")
                return ""
    
    def get_combined_text(self, processed_data: Dict) -> str:
        """
        Combine all extracted text into single string
        
        Args:
            processed_data: Output from process_folder()
            
        Returns:
            Combined text from all documents
        """
        combined_parts = []
        
        # Add PDF content
        for pdf in processed_data.get('pdfs', []):
            combined_parts.append(f"\n{'='*60}\nFILE: {pdf['filename']}\n{'='*60}\n")
            combined_parts.append(pdf['content'])
        
        # Add Word docs
        for doc in processed_data.get('word_docs', []):
            combined_parts.append(f"\n{'='*60}\nFILE: {doc['filename']}\n{'='*60}\n")
            combined_parts.append(doc['content'])
        
        # Add Excel files
        for excel in processed_data.get('excel_files', []):
            combined_parts.append(f"\n{'='*60}\nFILE: {excel['filename']}\n{'='*60}\n")
            if isinstance(excel['content'], dict):
                for sheet_text in excel['content'].get('text_summary', []):
                    combined_parts.append(sheet_text)
        
        # Add text files
        for txt in processed_data.get('text_files', []):
            combined_parts.append(f"\n{'='*60}\nFILE: {txt['filename']}\n{'='*60}\n")
            combined_parts.append(txt['content'])
        
        return "\n\n".join(combined_parts)
    
    def get_statistics(self, processed_data: Dict) -> Dict:
        """Get processing statistics"""
        return {
            'total_files': processed_data.get('total_files', 0),
            'pdf_files': len(processed_data.get('pdfs', [])),
            'excel_files': len(processed_data.get('excel_files', [])),
            'word_files': len(processed_data.get('word_docs', [])),
            'image_files': len(processed_data.get('images', [])),
            'text_files': len(processed_data.get('text_files', [])),
            'total_text_length': processed_data.get('total_text_length', 0),
            'errors': len(processed_data.get('errors', []))
        }


if __name__ == "__main__":
    # Test the module
    print("=" * 60)
    print("Document Processor Test")
    print("=" * 60)
    
    # Test with downloads folder
    downloads_path = Path(__file__).parent.parent / 'downloads'
    
    if downloads_path.exists():
        # Find first tender folder
        tender_folders = [f for f in downloads_path.iterdir() if f.is_dir()]
        
        if tender_folders:
            test_folder = tender_folders[2]
            print(f"\n📁 Testing with: {test_folder.name}\n")
            
            processor = DocumentProcessor()
            result = processor.process_folder(test_folder)
            
            stats = processor.get_statistics(result)
            print(f"\n📊 Statistics:")
            print(f"  Total Files: {stats['total_files']}")
            print(f"  PDFs: {stats['pdf_files']}")
            print(f"  Excel: {stats['excel_files']}")
            print(f"  Word: {stats['word_files']}")
            print(f"  Images: {stats['image_files']}")
            print(f"  Text: {stats['text_files']}")
            print(f"  Total Text: {stats['total_text_length']:,} characters")
            print(f"  Errors: {stats['errors']}")
            
            if result.get('errors'):
                print(f"\n⚠️ Errors:")
                for error in result['errors']:
                    print(f"  - {error}")
            
            # Show sample of extracted text
            combined_text = processor.get_combined_text(result)
            print(f"\n📄 Sample Text (first 500 chars):")
            print("-" * 60)
            print(combined_text[:500])
            print("-" * 60)
        else:
            print("No tender folders found in downloads/")
    else:
        print(f"Downloads folder not found: {downloads_path}")
